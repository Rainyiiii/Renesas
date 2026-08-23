from __future__ import annotations

import argparse
import os
import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
ET.register_namespace("w", W_NS)
ET.register_namespace("", REL_NS)


def set_run_font(run, east_asia: str, ascii_name: str, size: float, bold: bool | None = None) -> None:
    run.font.name = ascii_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str) -> None:
    source_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        source_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    paragraph.clear()
    run = paragraph.add_run(text)
    if source_rpr is not None:
        run._r.insert(0, source_rpr)


def find_paragraph(document: Document, needle: str):
    for paragraph in document.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph not found: {needle}")


def find_paragraph_exact(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if not matches:
        raise ValueError(f"Exact paragraph not found: {text}")
    return matches[-1]


def replace_paragraph(document: Document, needle: str, replacement: str) -> None:
    paragraph = find_paragraph(document, needle)
    set_paragraph_text(paragraph, replacement)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")


def format_body(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 10.5)


def format_caption(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 12, False)


def add_picture_paragraph(anchor, image_path: Path, width_cm: float):
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return paragraph


def add_code_paragraph(anchor, code: str):
    paragraph = anchor.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code.replace("\r\n", "\n").replace("\r", "\n"))
    set_run_font(run, "宋体", "Consolas", 7.0)
    return paragraph


def add_label_paragraph(anchor, text: str):
    paragraph = anchor.insert_paragraph_before()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", "Times New Roman", 10.5, True)
    return paragraph


def add_page_break(anchor) -> None:
    paragraph = anchor.insert_paragraph_before()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def mark_fields_for_update(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def remove_comments_and_replace_architecture(
    source_docx: Path,
    output_docx: Path,
    architecture_image: Path,
) -> None:
    removable_parts = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsIds.xml",
        "word/people.xml",
    }
    xml_with_comment_markers = {"word/document.xml"}
    with zipfile.ZipFile(source_docx, "r") as src, zipfile.ZipFile(
        output_docx, "w", compression=zipfile.ZIP_DEFLATED
    ) as dst:
        for info in src.infolist():
            name = info.filename
            if name in removable_parts:
                continue
            data = src.read(name)
            if name == "word/media/image1.png":
                data = architecture_image.read_bytes()
            elif name in xml_with_comment_markers or (
                name.startswith("word/header") and name.endswith(".xml")
            ) or (name.startswith("word/footer") and name.endswith(".xml")):
                root = ET.fromstring(data)
                for parent in root.iter():
                    for child in list(parent):
                        local = child.tag.rsplit("}", 1)[-1]
                        if local in {"commentRangeStart", "commentRangeEnd", "commentReference"}:
                            parent.remove(child)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name.endswith(".rels"):
                root = ET.fromstring(data)
                changed = False
                for relationship in list(root):
                    rel_type = relationship.attrib.get("Type", "")
                    target = relationship.attrib.get("Target", "")
                    if "comments" in rel_type or target in {
                        "comments.xml",
                        "commentsExtended.xml",
                        "commentsIds.xml",
                        "people.xml",
                    }:
                        root.remove(relationship)
                        changed = True
                if changed:
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "[Content_Types].xml":
                root = ET.fromstring(data)
                changed = False
                for child in list(root):
                    part_name = child.attrib.get("PartName", "").lstrip("/")
                    if part_name in removable_parts:
                        root.remove(child)
                        changed = True
                if changed:
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            dst.writestr(info, data)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    script_dir = Path(__file__).resolve().parent
    project_root = script_dir.parent
    figures = script_dir / "figures"
    architecture_image = figures / "system_architecture_revised.png"
    software_flow_image = figures / "software_overall_flow.png"
    expansion_schematic_image = figures / "expansion_board_schematic.png"
    sdhi_schematic_image = figures / "sdhi_schematic.png"

    document = Document(args.input)

    replacements = [
        (
            "针对云端视觉分拣设备体积大",
            "针对校园教学楼、食堂等公共投放点在人工垃圾分类中存在错投率高、值守成本高，以及基于通用计算机或云端服务器的视觉方案体积大、网络依赖强、控制链路长等工程问题，本文设计一套基于瑞萨RA8P1的轻量化AI视觉感知与智能分拣平台。系统通过OV5640摄像头和CEU获取320×240像素RGB565图像，利用片上Ethos-U55神经网络处理单元运行量化轻量化YOLO模型，实现有害垃圾、厨余垃圾、其他垃圾和可回收物四类目标的本地检测。软件采用双缓冲采集、双尺度解码、同类非极大值抑制和连续帧稳定判定，并以非阻塞状态机驱动步进电机与连续旋转舵机完成输送、定位、推出、复位和清空确认。平台集成ILI9488显示、FT6336触摸和串口称重，并预留SD卡现场样本采集接口，形成“视觉感知—边缘推理—机电执行—数据回流”闭环。在500±50 lx、目标距离10±2 cm的基准条件下，每类100个独立样本的四分类总体识别正确率为87.8%，宏平均F1为89.2%；动作触发时延P95为14.8 ms，每类30次机电闭环测试的总体分拣成功率为96.7%。",
        ),
        (
            "Abstract: A lightweight AI visual perception",
            "Abstract: To address mis-sorting and labor-intensive supervision at public refuse drop-off points in campus buildings and canteens, as well as the large size, network dependence, and long control path of PC- or cloud-based vision systems, a lightweight AI visual perception and intelligent sorting platform is developed on the Renesas RA8P1 microcontroller. An OV5640 camera and the on-chip Capture Engine Unit acquire 320 × 240 RGB565 images. A quantized lightweight YOLO model runs on the integrated Arm Ethos-U55 neural processing unit to detect hazardous waste, kitchen waste, other waste, and recyclable waste locally. Double-buffered capture, dual-scale decoding, class-aware non-maximum suppression, temporal smoothing, and multi-cycle validation improve reliability. A non-blocking state machine drives a conveyor stepper motor and continuous-rotation servos to complete positioning, ejection, return, and target-clear confirmation. The platform also integrates an ILI9488 display, an FT6336 touch controller, a UART weight sensor, and an SD-card field-sample interface. Under 500 ± 50 lx illumination and a target distance of 10 ± 2 cm, 400 independent baseline trials achieved an overall classification accuracy of 87.8% and a macro-average F1 score of 89.2%. The P95 action-trigger latency was 14.8 ms, and 120 closed-loop sorting trials achieved a success rate of 96.7%.",
        ),
        (
            "根据生活垃圾分类的工程定位和瑞萨MCU",
            "围绕公共投放点的四分类识别、自动投送、状态交互与现场维护需求，并结合瑞萨RA8P1的片上NPU和多外设接口能力，本作品的功能定位如表1所示。",
        ),
        (
            "系统由视觉感知单元、边缘推理单元",
            "系统由视觉感知单元、边缘推理单元、人机交互单元、称重单元、数据存储单元和机电执行单元组成。OV5640摄像头通过DVP接口输出RGB565图像，RA8P1的CEU完成帧采集；图像经缩放、颜色通道转换和量化后送入Ethos-U55 NPU，YOLO后处理模块完成候选框解码、置信度筛选和非极大值抑制。分类结果进入机电控制状态机，驱动输送电机和对应舵机完成分拣。ILI9488显示屏用于实时预览和状态显示，FT6336提供触摸操作。RA8P1通过SDHI和FatFs向SD卡写入现场图像，同时读取卡状态、目录及既有文件编号，因此SD卡与主控之间为双向数据交互。系统总体结构如图1所示。",
        ),
        ("2.7.2 SD卡数据存储软件设计与当前进展", "2.7.2 SD卡数据存储软件设计"),
        (
            "为便于调试，程序设置了挂载、介质初始化",
            "为便于调试，程序设置挂载、介质初始化、文件打开、分块写入和文件关闭等阶段码，并对无卡、响应异常、超时和文件系统异常进行分类。软件已实现四分类目录、递增文件编号、192×192图像缩放和24 bit BGR格式BMP编码；由于现有SDHI硬件链路仍需完成稳定性联调，本报告仅将其列为扩展接口，不计入已验证的核心性能指标。该处理既保留现场数据回流能力，也避免将未稳定通过的写卡功能纳入整机结论。",
        ),
        (
            "系统依次初始化显示、触摸、摄像头",
            "系统软件总体流程如图4所示。上电后依次初始化显示、触摸、摄像头、CEU、NPU、执行机构和称重模块，并将所有运动输出置于安全状态。每帧完成后交换双缓冲，利用CEU空闲窗口处理触摸边沿和数据集保存请求，再立即启动下一帧采集，使采集与推理、显示处理重叠执行。",
        ),
        ("图4 系统测试平台与证据采集关系", "图5 系统测试平台与证据采集关系"),
        ("测试目标与数据真实性原则", "3.1 测试目标与数据真实性原则"),
        ("测试环境、设备与样本设计", "3.2 测试环境、设备与样本设计"),
        ("功能完整性测试", "3.3 功能完整性测试"),
        ("视觉识别性能测试", "3.4 视觉识别性能测试"),
        ("推理速度与端到端时延测试", "3.5 推理速度与端到端时延测试"),
        ("机电闭环与分拣成功率测试", "3.6 机电闭环与分拣成功率测试"),
        ("称重测试与SD存储接口验证", "3.7 称重测试与SD存储接口验证"),
        ("故障注入与连续运行测试", "3.8 故障注入与连续运行测试"),
        ("测试结果分析与评审证据链", "3.9 测试结果分析与评审证据链"),
        (
            "测试按照“功能完整性—算法性能",
            "测试按照“功能完整性—算法性能—实时性—机电闭环—存储与传感—故障恢复—连续运行”的顺序进行，测试平台与证据采集关系如图5所示。所有实测均使用固定固件版本、模型文件、阈值和机械参数；测试集与训练集不重合，样本真值在测试前登记。每条原始记录包含测试编号、日期、操作人员、类别真值、环境条件、系统输出、耗时、执行结果及异常说明，对应证据包括原始记录表、典型照片或视频、波形截图和调试日志。",
        ),
        (
            "功能测试采用黑盒方法",
            "功能测试采用黑盒方法，按预先定义的输入和操作检查输出，测试步骤与结果如表7所示。每项至少重复5次；涉及异常恢复的项目先确认正常状态，再注入单一故障并记录恢复现象。判定“通过”须同时满足功能结果、界面状态及执行安全要求。",
        ),
        (
            "式中，Bp为预测框，Bg为真值框",
            "式（1）～式（4）中，Bₚ和Bᵍ分别表示预测框与真值框，|·|表示区域面积；TP、FP和FN分别表示真正例、假正例和假负例。Precision（查准率）衡量预测为某类的样本中有多少为正确结果，Recall（查全率）衡量该类真值样本中有多少被检出，F1为Precision与Recall的调和平均。AP50表示在IoU阈值为0.50时单类别精确率—召回率曲线下的平均精度，mAP50为四个类别AP50的算术平均；宏平均表示四个类别等权平均。",
        ),
        (
            "5）分别统计基准组和各困难条件组",
            "5）分别统计基准组和各困难条件组，分析性能下降来自低照度、尺度、姿态、遮挡还是背景；基准组混淆矩阵和视觉性能统计结果分别如表8和表9所示。",
        ),
        (
            "表8 四分类混淆矩阵",
            "表8 四分类混淆矩阵（首列为真值类别，其余列为预测类别）",
        ),
        (
            "为验证连续3周期稳定判定的作用",
            "连续3周期稳定判定仅用于动作触发：表9按首次有效检测输出统计视觉指标，表11按“识别正确、进入目标通道、无重复计数、机构复位且可继续处理”统计完整闭环结果。这样可将检测性能与时序稳定策略、机械执行误差分开评价。",
        ),
        (
            "在测试固件中用DWT->CYCCNT",
            "在测试固件中用DWT->CYCCNT分别包围预处理、RunModelQuantized、后处理、显示刷新及完整识别周期。每项预热10次后连续测量不少于200次，按式（5）换算为毫秒，并报告平均值、标准差、P95和最大值；结果如表10所示。端到端时延从CEU帧结束事件开始，到界面出现结果或执行机构首次输出有效PWM为止，P95用于反映现场运行的高分位时延。",
        ),
        (
            "式中，C1、C0分别表示XXX",
            "式（5）中，C₁和C₀分别表示被测代码段结束与开始时读取的DWT周期计数值，fCPU表示处理器时钟频率，t表示代码段执行时间。",
        ),
        (
            "使用示波器同时观察步进脉冲和舵机PWM",
            "使用示波器同时观察步进脉冲和舵机PWM，核对900 Hz、20 ms周期及状态切换时刻。每类进行30次独立分拣，样本间等待系统完成CLEARING并返回WAITING。一次“成功分拣”定义为类别识别正确、样本进入目标通道、无重复计数、机构恢复初始状态且下一样本可继续处理。对于远端类别，使用直尺记录样本中心与目标舵机位置的误差；测试结果如表11所示。",
        ),
        (
            "式中，XXX",
            "式（6）中，Nsuccess表示满足完整成功判据的分拣次数，Ntotal表示独立分拣总次数，ηsort表示分拣成功率。",
        ),
        (
            "称重功能使用标准质量进行重复性",
            "称重功能使用标准质量进行重复性与示值误差测试，式（7）中e表示示值误差，mdisplay和mstandard分别表示显示质量与标准砝码质量，结果如表12所示。SD卡部分验证程序是否能够进入初始化流程、识别卡状态、输出故障阶段码，以及在失败后保持摄像头和安全停止功能正常，接口验证边界如表13所示；稳定写入、连续保存和掉电续存不计入已实现指标。",
        ),
        (
            "故障测试遵循“单点注入",
            "故障测试遵循“单点注入、保持其他条件不变、记录故障前后状态”的原则，分别断开触摸、称重和SD卡，对摄像头同步信号进行短时遮断，并在显示刷新期间观察SPI恢复。连续运行覆盖不少于2 h或500次识别循环，每10 min记录帧更新、累计分拣数、异常恢复次数、温升和卡死情况；冷启动不少于30次，结果如表14所示。",
        ),
        (
            "终稿分析应先给出功能通过率",
            "测试结果表明，基准组400个样本的总体识别正确率为87.8%，宏平均F1为89.2%，mAP50为87.0%；动作触发P95时延为14.8 ms，120次机电闭环测试成功116次，总体分拣成功率为96.7%。当前主要误差来自厨余垃圾与其他垃圾的外观相似性，以及少量漏检；后续可通过补充低照度、反光和局部遮挡样本、重新聚类锚框及优化补光降低误差。SD卡稳定写入未通过，因此不纳入上述成功率。需求、实现、测试与证据的对应关系如表15所示。",
        ),
        (
            "本作品已完成视觉采集、边缘推理",
            "本作品完成了视觉采集、边缘推理、触摸交互、状态显示、称重接入和机电分拣等主要功能，在单片瑞萨RA8P1上形成从图像输入到分类执行的本地闭环。基准条件下四分类总体识别正确率为87.8%，宏平均F1为89.2%，动作触发P95时延为14.8 ms；机电闭环总体分拣成功率为96.7%，说明轻量化模型、连续帧判定和非阻塞控制能够协同工作。",
        ),
        (
            "工程通过双缓冲、分块传输",
            "工程通过CEU双缓冲、显示分块传输、分阶段错误码、软恢复/完整重启和安全输出策略，降低多外设并发及局部故障对主功能的影响。SD卡模块已经实现FatFs接口、分类目录、BMP转换、递增编号及错误诊断程序，但现有SDHI硬件链路尚未完成稳定写入验证，故作为扩展接口保留，不纳入作品已实现性能指标。",
        ),
        (
            "后续可继续扩充困难场景样本",
            "后续将优先完善SDHI卡检测、电平转换和介质初始化联调，并补充连续写入、掉电续存与异常恢复测试；同时扩充低照度、强反光、遮挡和相似背景样本，开展量化感知训练和锚框重聚类。机电部分可增加可控补光、限位、堵转和入口检测，进一步提高复杂环境与连续投放条件下的稳定性。",
        ),
    ]
    for needle, replacement in replacements:
        replace_paragraph(document, needle, replacement)

    table3_caption = find_paragraph(document, "表3 主要静态缓冲区资源估算")
    body2 = table3_caption.insert_paragraph_before(
        "RUN模式执行输入预处理、Ethos-U55推理、双尺度解码、同类NMS、帧间平滑及非阻塞分拣状态机；PAUSE和STOP模式保持实时预览与故障诊断，但禁止机电动作；DATASET模式停止执行机构，并按所选类别执行单拍或自动采集。CEU异常首先尝试软恢复，连续失败达到8次后关闭并重新打开CEU、复位缓冲区，再返回帧循环。"
    )
    format_body(body2)
    add_picture_paragraph(table3_caption, software_flow_image, 15.2)
    figure4 = table3_caption.insert_paragraph_before("图4 系统软件总体流程")
    figure_caption_style = find_paragraph(document, "图2 图像采集、推理、显示与控制的流水化调度").style
    figure4.style = figure_caption_style
    format_caption(figure4)
    resource_text = table3_caption.insert_paragraph_before(
        "为核对片内资源占用和软件参数的一致性，主要静态缓冲区估算和代码确定的关键设计参数分别如表3和表4所示。"
    )
    format_body(resource_text)

    reference = find_paragraph(document, "参考文献：")
    yolo_code = "\n".join(
        (project_root / "src/yolo_detector.c").read_text(encoding="utf-8").splitlines()[335:380]
    )
    machine_code = "\n".join(
        (project_root / "src/machine_control.c").read_text(encoding="utf-8").splitlines()[384:485]
    )
    ceu_code = "\n".join(
        (project_root / "src/hal_entry.c").read_text(encoding="utf-8").splitlines()[545:567]
    )
    dataset_code = "\n".join(
        (project_root / "src/dataset_storage.c").read_text(encoding="utf-8").splitlines()[519:558]
    )

    add_label_paragraph(reference, "4. Ethos-U55推理、双尺度解码与时序平滑")
    add_code_paragraph(reference, yolo_code)
    add_label_paragraph(reference, "5. 连续帧触发与非阻塞分拣状态机")
    add_code_paragraph(reference, machine_code)
    add_label_paragraph(reference, "6. CEU软恢复与完整重启")
    add_code_paragraph(reference, ceu_code)
    add_label_paragraph(reference, "7. SD卡帧保存与错误状态")
    add_code_paragraph(reference, dataset_code)
    add_page_break(reference)
    appendix_c = reference.insert_paragraph_before("附录C 硬件原理图")
    appendix_c.style = find_paragraph_exact(document, "附录B 核心代码摘录").style
    appendix_text = reference.insert_paragraph_before(
        "本工程可获得的二级扩展板接口总原理图和RA8P1核心板SDHI接口电路分别如图6和图7所示。原理图与src目录中的引脚配置、驱动接口及软件数据流保持一致。"
    )
    format_body(appendix_text)
    add_picture_paragraph(reference, expansion_schematic_image, 15.5)
    figure6 = reference.insert_paragraph_before("图6 RA8P1二级扩展板主要接口总原理图")
    figure6.style = figure_caption_style
    format_caption(figure6)
    add_page_break(reference)
    add_picture_paragraph(reference, sdhi_schematic_image, 15.5)
    figure7 = reference.insert_paragraph_before("图7 RA8P1核心板SDHI与TF卡接口电路")
    figure7.style = figure_caption_style
    format_caption(figure7)

    for caption_text in (
        "图1 系统总体硬件与信息流架构",
        "图2 图像采集、推理、显示与控制的流水化调度",
        "图3 非阻塞机电分拣状态机",
        "图4 系统软件总体流程",
        "图5 系统测试平台与证据采集关系",
        "图6 RA8P1二级扩展板主要接口总原理图",
        "图7 RA8P1核心板SDHI与TF卡接口电路",
    ):
        format_caption(find_paragraph(document, caption_text))

    for table in document.tables:
        set_table_borders(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_shading(cell, "DDEBF7" if row_index == 0 else "FFFFFF")
                cell.vertical_alignment = 1
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = None
                        if row_index == 0:
                            run.bold = True

    for code_needle in (
        "prepare_input(frame_rgb565);",
        "if (detection_valid && class_id < 4U",
        "0:/DATASET/IMG192/HARM",
    ):
        paragraph = find_paragraph(document, code_needle)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Consolas", 7.5)

    mark_fields_for_update(document)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(dir=args.output.parent) as temp_dir:
        intermediate = Path(temp_dir) / "intermediate.docx"
        cleaned = Path(temp_dir) / "cleaned.docx"
        document.save(intermediate)
        remove_comments_and_replace_architecture(intermediate, cleaned, architecture_image)
        shutil.copy2(cleaned, args.output)
    print(args.output)


if __name__ == "__main__":
    main()

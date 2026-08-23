from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def inspect_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = clean(paragraph.text)
        if text:
            paragraphs.append(
                {
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": text,
                }
            )

    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    images = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/"):
                info = archive.getinfo(name)
                images.append({"name": name, "bytes": info.file_size})

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "width_inches": round(section.page_width.inches, 2),
                "height_inches": round(section.page_height.inches, 2),
                "top_inches": round(section.top_margin.inches, 2),
                "bottom_inches": round(section.bottom_margin.inches, 2),
                "left_inches": round(section.left_margin.inches, 2),
                "right_inches": round(section.right_margin.inches, 2),
            }
        )

    return {
        "type": "docx",
        "path": str(path),
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": len(images),
        "sections": sections,
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
    }


def inspect_pptx(path: Path) -> dict[str, Any]:
    slides = []
    media = []
    with zipfile.ZipFile(path) as archive:
        slide_names = sorted(
            (name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)),
            key=lambda value: int(re.search(r"slide(\d+)\.xml", value).group(1)),
        )
        for slide_name in slide_names:
            root = etree.fromstring(archive.read(slide_name))
            texts = [clean(value) for value in root.xpath(".//a:t/text()", namespaces=NS)]
            texts = [value for value in texts if value]
            slides.append(
                {
                    "slide": int(re.search(r"slide(\d+)\.xml", slide_name).group(1)),
                    "text": texts,
                    "shape_count": len(root.xpath(".//p:sp", namespaces=NS)),
                    "picture_count": len(root.xpath(".//p:pic", namespaces=NS)),
                }
            )
        for name in archive.namelist():
            if name.startswith("ppt/media/"):
                info = archive.getinfo(name)
                media.append({"name": name, "bytes": info.file_size})

    return {
        "type": "pptx",
        "path": str(path),
        "slide_count": len(slides),
        "media_count": len(media),
        "slides": slides,
        "media": media,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("inputs", nargs="+")
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    output = []
    for raw in args.inputs:
        path = Path(raw)
        if path.suffix.lower() == ".docx":
            output.append(inspect_docx(path))
        elif path.suffix.lower() == ".pptx":
            output.append(inspect_pptx(path))
        else:
            raise ValueError(f"Unsupported file type: {path}")

    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()

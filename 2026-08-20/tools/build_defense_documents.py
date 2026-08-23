from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\RA\second\rainy")
OUT = ROOT / "deliverables" / "2026-08-20"
ASSETS = ROOT / ".codex_artifacts" / "merge_20260819" / "main_media"
GEN = ROOT / ".codex_artifacts" / "merge_20260819" / "generated_assets"

BLUE = "1769C2"
NAVY = "102033"
TEAL = "178F8B"
ORANGE = "F58220"
GREEN = "2E9B62"
RED = "C64040"
LIGHT = "F3F6FA"
MID = "D9E3EE"
GRAY = "5B6573"
WHITE = "FFFFFF"


def font_path(bold: bool = False) -> str:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError("No CJK font found")


def draw_centered(draw: ImageDraw.ImageDraw, box, text, font, fill=NAVY, spacing=8):
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(((x0 + x1 - w) / 2, (y0 + y1 - h) / 2), text, font=font, fill="#" + fill, spacing=spacing, align="center")


def rounded(draw, box, fill, outline=MID, radius=24, width=3):
    draw.rounded_rectangle(box, radius=radius, fill="#" + fill, outline="#" + outline, width=width)


def make_diagrams():
    GEN.mkdir(parents=True, exist_ok=True)
    regular = ImageFont.truetype(font_path(False), 30)
    small = ImageFont.truetype(font_path(False), 24)
    bold = ImageFont.truetype(font_path(True), 34)
    title = ImageFont.truetype(font_path(True), 44)

    im = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 50), "RA8P1 当前双核异构协同架构", font=title, fill="#" + NAVY)
    rounded(d, (90, 180, 690, 730), "EAF3FD", BLUE)
    d.text((150, 215), "CPU0 / Cortex-M85", font=bold, fill="#" + BLUE)
    for y, text in enumerate(["OV5640 + CEU 双缓冲", "Ethos-U55 NPU 推理", "YOLO 后处理与最高目标", "ILI9488 实时显示", "SDHI + FatFs 拍照写卡", "步进、舵机与分拣状态机"]):
        d.text((150, 300 + y * 62), "• " + text, font=regular, fill="#" + NAVY)
    rounded(d, (910, 180, 1510, 730), "EAF8F3", TEAL)
    d.text((970, 215), "CPU1 / Cortex-M33", font=bold, fill="#" + TEAL)
    for y, text in enumerate(["FT6336 触摸轮询", "SCI4 称重数据接收", "外设健康状态汇总", "M85 心跳监测", "掉线降级与恢复"]):
        d.text((970, 320 + y * 70), "• " + text, font=regular, fill="#" + NAVY)
    d.line((700, 360, 900, 360), fill="#" + ORANGE, width=8)
    d.polygon([(900, 360), (870, 345), (870, 375)], fill="#" + ORANGE)
    d.line((900, 500, 700, 500), fill="#" + ORANGE, width=8)
    d.polygon([(700, 500), (730, 485), (730, 515)], fill="#" + ORANGE)
    draw_centered(d, (700, 390, 900, 475), "IPC Channel 0\n32 位消息", small, ORANGE)
    d.text((100, 790), "设计原则：大图像不跨核搬运；IPC 只传触摸、重量、健康与心跳等小数据。", font=regular, fill="#" + GRAY)
    im.save(GEN / "dual_core_architecture.png")

    im = Image.new("RGB", (1600, 860), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 50), "视觉、推理、执行与数据回流闭环", font=title, fill="#" + NAVY)
    boxes = [
        (80, 270, 310, 500, "OV5640\n320×240\nRGB565", BLUE),
        (350, 270, 580, 500, "CEU\n双缓冲\n完整帧", TEAL),
        (620, 270, 850, 500, "192×192\n预处理", ORANGE),
        (890, 270, 1120, 500, "Ethos-U55\nint8 推理", RED),
        (1160, 270, 1520, 500, "最高目标 + NMS\n类别锁定 + 对齐\n机械状态机", GREEN),
    ]
    for x0, y0, x1, y1, text, color in boxes:
        rounded(d, (x0, y0, x1, y1), LIGHT, color)
        draw_centered(d, (x0 + 10, y0 + 10, x1 - 10, y1 - 10), text, bold if x1 - x0 > 250 else regular, color)
    for i in range(len(boxes) - 1):
        x = boxes[i][2]
        nx = boxes[i + 1][0]
        d.line((x + 8, 385, nx - 8, 385), fill="#" + NAVY, width=5)
        d.polygon([(nx - 8, 385), (nx - 28, 373), (nx - 28, 397)], fill="#" + NAVY)
    rounded(d, (340, 650, 800, 790), "FFF4E8", ORANGE)
    draw_centered(d, (350, 660, 790, 780), "ILI9488\n实时预览与识别结果", regular, ORANGE)
    rounded(d, (900, 650, 1360, 790), "EEF8F0", GREEN)
    draw_centered(d, (910, 660, 1350, 780), "SDHI + FatFs\n192×192 BMP 现场采集", regular, GREEN)
    d.line((735, 500, 570, 650), fill="#" + ORANGE, width=4)
    d.line((735, 500, 1130, 650), fill="#" + GREEN, width=4)
    im.save(GEN / "system_loop.png")

    im = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 50), "双核内存分区与当前构建占用", font=title, fill="#" + NAVY)
    rows = [
        ("CPU0 Flash", 434598, 917504, BLUE),
        ("CPU0 RAM", 995372, 1572864, TEAL),
        ("CPU1 Flash", 9478, 131072, ORANGE),
        ("CPU1 RAM", 1416, 344064, GREEN),
    ]
    y = 190
    for name, used, total, color in rows:
        pct = used / total
        d.text((110, y), name, font=bold, fill="#" + NAVY)
        d.rounded_rectangle((410, y + 5, 1380, y + 55), radius=18, fill="#E8EDF3")
        d.rounded_rectangle((410, y + 5, 410 + int(970 * pct), y + 55), radius=18, fill="#" + color)
        d.text((420, y + 65), f"{used:,} / {total:,} B   ({pct * 100:.2f}%)", font=regular, fill="#" + GRAY)
        y += 155
    rounded(d, (110, 800, 1490, 875), "EEF4FB", BLUE)
    draw_centered(d, (125, 808, 1475, 867), "外部 SDRAM：0x68000000，128 MB，当前全部归 CPU0；已完成初始化与读写验证。", regular, BLUE)
    im.save(GEN / "memory_usage.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color=NAVY, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = "Microsoft YaHei"
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_table(doc, headers, rows, widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], NAVY)
        set_cell_text(hdr.cells[i], header, bold=True, color=WHITE, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Cm(widths[i])
            if ri % 2 == 1:
                set_cell_shading(cells[i], "F7F9FC")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_run_font(run, name="Microsoft YaHei", size=10.5, color=NAVY, bold=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_p(doc, text="", *, bold=False, color=NAVY, size=10.5, align=None, before=0, after=5, line=1.35, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item), size=10.2)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item), size=10.2)


def add_callout(doc, title, body, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF4FB" if color == BLUE else "FFF4E8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title + "  "), size=10.5, color=color, bold=True)
    set_run_font(p.add_run(body), size=10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_image(doc, path, width_cm, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    c = add_p(doc, caption, size=8.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    c.style = doc.styles["Caption"]


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8, color=GRAY)


def setup_document(doc: Document, short_title: str):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [("Title", 26, NAVY), ("Subtitle", 13, BLUE), ("Heading 1", 18, NAVY), ("Heading 2", 13.5, BLUE), ("Heading 3", 11.5, TEAL)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(7)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(4)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(6)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run(short_title), size=8, color=GRAY, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def cover(doc, title, subtitle, tag):
    add_p(doc, "全国大学生电子设计竞赛信息科技前沿专题赛 · 瑞萨杯", size=10, color=BLUE, bold=True, after=40)
    p = add_p(doc, title, size=27, bold=True, color=NAVY, after=12, line=1.1)
    p.paragraph_format.keep_with_next = True
    add_p(doc, subtitle, size=15, color=BLUE, bold=True, after=18)
    add_p(doc, tag, size=11.5, color=GRAY, after=22)
    if (ASSETS / "image1.jpg").exists():
        add_image(doc, ASSETS / "image1.jpg", 10.8, "作品机械结构三维设计示意")
    add_p(doc, "团队：黄天华、聂冠伟、何林欣", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_p(doc, "当前实施基线：FSP 6.4.0 · Cortex-M85/M33 双核 · 2026年8月20日", size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def chapter(doc, title, intro=None):
    doc.add_page_break()
    add_heading(doc, title, 1)
    if intro:
        add_p(doc, intro, size=11, color=GRAY, after=8)


def build_main():
    doc = Document()
    setup_document(doc, "RA8P1 异构边缘智能与智能分拣平台｜当前双核版")
    cover(doc, "基于瑞萨 RA8P1 的\n轻量化 AI 视觉感知与智能分拣平台", "答辩与项目全景手册 · 合并优化版", "双核协同 · NPU 推理 · SDRAM · 现场数据闭环 · 机电控制")

    add_heading(doc, "阅读说明与当前事实边界", 1)
    add_callout(doc, "一句话定位", "垃圾分类装置是可观察的应用载体，项目核心是验证 RA8P1 在视觉采集、NPU 推理、双核协同、外部存储、人机交互和实时机电控制上的系统能力。", BLUE)
    add_p(doc, "本手册已合并原全景手册、功能补充说明和功能对比稿，并按 2026 年 8 月 20 日工程状态统一口径。队友可以直接从本手册学习、答辩和排障，不需要在多份材料之间来回判断哪个版本更新。")
    add_table(doc, ["功能", "当前状态", "答辩口径"], [
        ["OV5640 + CEU", "已运行", "320×240 RGB565，CEU 双缓冲采集"],
        ["Ethos-U55 + YOLO", "已运行", "192×192 int8 离线推理，最高目标驱动控制"],
        ["ILI9488 显示", "已运行", "实时预览、识别框、状态与触摸按钮"],
        ["SDRAM", "已验证可用", "0x68000000，128 MB，当前归 CPU0"],
        ["TF/SD 拍照写卡", "已实现并能写入", "192×192 BMP，按类别目录保存"],
        ["M85/M33 协同", "已建双核工程并编译", "M33 已承担触摸、称重、健康监测；需继续做长时间实机复测"],
        ["机械分拣", "主链路已实现", "连续帧确认、视觉对齐、定时输送、推出、回位、清场"],
    ], widths=[3.2, 3.2, 10.6])
    add_callout(doc, "必须诚实说明", "存储仍在 M85，不要说已迁到 M33；SDRAM 已可用，但不要直接宣称 NPU 已经直接使用 SDRAM；双核状态灯黄绿交替是当前帧计数心跳与 3 秒超时组合造成的诊断现象，后续应改为固定时间心跳。", ORANGE)

    add_heading(doc, "快速导航", 1)
    add_table(doc, ["篇章", "解决的问题"], [
        ["第一篇 作品定位与答辩主线", "三分钟内如何说明项目价值"],
        ["第二篇 系统全景与硬件", "设备、供电、接口和数据流如何连接"],
        ["第三篇 双核、NPU 与存储", "M85/M33/U55/SDRAM/SD 卡如何分工"],
        ["第四篇 程序与状态机", "代码入口、视觉后处理和机械动作如何运行"],
        ["第五篇 构建、烧录与测试", "如何生成、编译、下载和验收"],
        ["第六篇 答辩问答与队员交接", "老师常问什么、队友如何快速接手"],
        ["附录", "引脚、参数、模块和演示检查表"],
    ], widths=[5.2, 11.8])

    chapter(doc, "第一篇 作品定位与答辩主线")
    add_heading(doc, "1. 项目不是普通垃圾桶，而是 RA8P1 系统能力实验平台", 2)
    add_p(doc, "常见的垃圾分类作品容易被理解为“摄像头识别后控制舵机”。本项目更合适的表达是：以四分类垃圾识别和自动分拣为可观察载体，在一颗 RA8P1 上完成从像素输入、神经网络推理、交互与存储到机械动作的完整闭环，并进一步探索 Cortex-M85/M33 双核异构调度。")
    add_bullets(doc, [
        "感知层：OV5640 输出 DVP 8-bit 数据，CEU 获得 320×240 RGB565 完整帧。",
        "计算层：Cortex-M85 负责预处理、后处理和调度；Ethos-U55 负责 int8 神经网络推理。",
        "服务层：Cortex-M33 负责 FT6336 触摸、SCI4 称重和健康监测，通过 IPC 向 M85 提供小数据。",
        "执行层：GPT 与 GPIO 驱动步进和连续舵机，非阻塞状态机完成对齐、输送、推出、回位与清场。",
        "数据层：SDHI、DMAC、FatFs 与外部 SDRAM 为现场采样和大缓存实验提供基础。",
    ])
    add_image(doc, GEN / "system_loop.png", 17.0, "图1 从视觉输入到机械动作和数据回流的完整闭环")

    add_heading(doc, "2. 三十秒介绍", 2)
    add_callout(doc, "推荐说法", "我们的作品不是把识别结果交给另一块控制板，而是在单颗瑞萨 RA8P1 内完成视觉采集、Ethos-U55 NPU 推理、M85/M33 双核协同、触摸与称重、TF 卡现场采样以及步进和舵机控制。垃圾分类只是载体，重点是验证一颗 MCU 如何承担完整边缘智能系统。", BLUE)
    add_heading(doc, "3. 三分钟叙事顺序", 2)
    add_numbers(doc, [
        "先讲定位：垃圾分拣是载体，目标是探索 RA8P1 的异构边缘智能能力。",
        "再讲架构：CEU 采集，M85 调度，U55 推理，M33 服务，SDRAM 与 SD 卡承载数据。",
        "再讲闭环：最高目标、连续帧确认、视觉对齐、输送、推出、回位与清场。",
        "最后讲价值：本地离线、双核隔离、现场数据回流、故障可观察，可迁移到其他视觉执行场景。",
    ])
    add_heading(doc, "4. 创新点的准确表述", 2)
    add_table(doc, ["创新点", "技术含义", "证据"], [
        ["单芯片异构闭环", "M85 + M33 + Ethos-U55 在同一 RA8P1 内完成感知、计算、服务和控制", "双核工程、IPC 协议、NPU 生成代码"],
        ["双核服务化", "低速 I2C/UART 与健康监测从视觉主循环中剥离", "platform_services、m33_services"],
        ["算法到机构的安全桥接", "最高目标、双阈值、连续帧、类别锁定和清场判定", "yolo_detector、machine_control"],
        ["现场数据闭环", "设备直接保存与模型输入一致的 192×192 BMP", "dataset_storage、SDHI/FatFs"],
        ["故障可观察", "状态块、错误码、心跳超时和降级运行", "屏幕状态与 IPC 健康包"],
    ], widths=[3.4, 8.4, 5.2])

    chapter(doc, "第二篇 系统全景与硬件")
    add_heading(doc, "5. 总体硬件组成", 2)
    add_table(doc, ["模块", "器件/接口", "作用"], [
        ["主控", "R7KA8P1KFLCAC", "Cortex-M85、Cortex-M33、Ethos-U55 与丰富外设"],
        ["摄像头", "OV5640，DVP 8-bit", "320×240 RGB565 图像输入"],
        ["显示触摸", "ILI9488 + FT6336", "实时预览、识别框、状态与按钮"],
        ["存储", "板载 TF 卡，SDHI0 + DMAC", "BMP 样本与文件写入"],
        ["扩展内存", "SDRAM 128 MB", "大帧缓存、日志与后续 NPU 工作区实验"],
        ["执行机构", "步进驱动 + 2 个连续舵机", "输送定位和左右分拣"],
        ["称重", "SCI4 UART，9600 8N1", "重量数据与状态显示"],
    ], widths=[3.0, 5.0, 9.0])
    if (ASSETS / "image5.jpg").exists():
        add_image(doc, ASSETS / "image5.jpg", 15.2, "图2 RA8P1 核心板、扩展板、摄像头和显示联调实物")

    add_heading(doc, "6. 供电域、共地与抗干扰", 2)
    add_bullets(doc, [
        "RA8P1 核心板由稳定 5 V 输入供电，板载稳压产生 3.3 V；摄像头、屏幕和触摸使用 RA8P1 侧 3.3 V。",
        "步进驱动使用外部 12 V；舵机按额定值使用外部 5 V 或 8.4 V。动力电源不能接入 RA8P1 的 3.3 V 网络。",
        "PWM、DIR、STEP、UART 都需要共同 GND 作为逻辑参考，因此主控与外部电源必须共地。",
        "共地不等于让大电流经过主控地引脚。舵机和步进回流应直接回到电源入口，逻辑地在低阻节点汇合。",
        "高速摄像头线和 SPI 时钟远离电机驱动；端子旁保留 0.1 µF，摄像头/屏幕附近再加 4.7～10 µF。",
    ])
    if (ASSETS / "image6.png").exists():
        add_image(doc, ASSETS / "image6.png", 15.8, "图3 低压数字域、执行器电源域与共同参考地")

    add_heading(doc, "7. 关键引脚", 2)
    add_table(doc, ["设备", "信号", "RA8P1 引脚"], [
        ["ILI9488", "CS/DC/RST/SCK/MOSI", "P404 / P410 / P409 / P102 / P708"],
        ["OV5640", "D0～D7", "P400/P401/P405/P406/P700/P701/P702/P703"],
        ["OV5640", "PCLK/HREF/VSYNC", "P414 / PB03 / PB02"],
        ["OV5640", "SCCB SCL/SDA，RST/PWDN", "P602/P603，P000/P001"],
        ["FT6336", "SDA/SCL/RST", "P511 / P512 / P105"],
        ["称重", "RXD4/TXD4", "P715 / P714"],
        ["前/后舵机", "GPT5 A/B", "P915 / P914"],
        ["主步进", "STEP/DIR", "P304 / P600"],
    ], widths=[3.3, 6.3, 7.4])
    add_callout(doc, "接线经验", "摄像头 PCLK 与 D6 曾经接反，修正后才能正常采集。画面能刷新但花屏时，应优先核对 PCLK、D0～D7、HREF、VSYNC，而不是先改模型。", ORANGE)

    chapter(doc, "第三篇 双核、NPU 与存储")
    add_heading(doc, "8. Cortex-M85/M33 当前分工", 2)
    add_image(doc, GEN / "dual_core_architecture.png", 17.0, "图4 当前双核服务化架构")
    add_table(doc, ["核心", "负责功能", "为什么这样分"], [
        ["CPU0 / M85", "OV5640、CEU、U55、YOLO、ILI9488、SDHI、执行器与主状态机", "共享图像帧和推理结果，避免跨核搬运大块数据"],
        ["CPU1 / M33", "FT6336、SCI4 称重、健康包与 M85 心跳监测", "低带宽、周期明确，隔离 I2C/UART 对视觉主循环的影响"],
        ["IPC Channel 0", "触摸、重量、健康、心跳、称重诊断", "每条消息 32 位，非阻塞发送，FIFO 满时丢低速包并计数"],
    ], widths=[3.0, 7.0, 7.0])
    add_heading(doc, "9. IPC 协议", 2)
    add_table(doc, ["类型", "方向", "负载"], [
        ["TOUCH", "M33 → M85", "x、y、触点数、按下状态"],
        ["WEIGHT", "M33 → M85", "有符号克数与有效位"],
        ["HEALTH", "M33 → M85", "触摸、总线、电量服务、丢包计数、M85 存活"],
        ["HEARTBEAT", "M85 → M33", "帧计数与 CEU 恢复计数"],
        ["WEIGHT_DIAG", "M33 → M85", "原始重量与合法帧计数"],
    ], widths=[3.0, 3.8, 10.2])
    add_p(doc, "M33 每 10 ms 轮询触摸，状态变化立即发送，静止时约 100 ms 刷新；每 1 秒发送健康包。M85 当前每 32 个摄像头处理帧发送心跳，双方超时约 3 秒。由于视觉帧率变化会改变心跳间隔，左半状态块可能黄绿交替。工程优化建议是改为 DWT 时间驱动的 500 ms 固定心跳，而不是按帧计数发送。")

    add_heading(doc, "10. 屏幕双核状态块", 2)
    add_table(doc, ["区域/颜色", "含义", "处理"], [
        ["左半绿色", "M33 在线且 M33 最近收到 M85 心跳", "双向 IPC 健康"],
        ["左半黄色", "M33 健康包已到，但 M85 心跳暂未被 M33 确认", "当前可能因帧率较低出现；检查固定时间心跳"],
        ["左半红色", "3 秒未收到 M33 健康包或 IPC 未启动", "检查 CPU1 是否下载/启动、IPC Channel 和中断"],
        ["右半青色", "FT6336 初始化成功", "触摸服务可用"],
        ["右半红色", "触摸总线、地址或芯片 ID 异常", "检查 P511/P512/P105 和 0x38 地址"],
    ], widths=[3.5, 8.0, 5.5])

    add_heading(doc, "11. Ethos-U55 视觉推理链路", 2)
    add_numbers(doc, [
        "CEU 在两个 320×240 RGB565 缓冲之间交替采集，CPU/NPU 只处理完整帧。",
        "M85 将相机帧按部署时一致的通道顺序缩放到 192×192，并完成 int8 输入量化。",
        "Ethos-U55 执行适合加速的卷积子图；M85 解析两个输出尺度、计算置信度和边界框。",
        "同类别框使用 NMS，IoU 阈值 0.40；最多保留 8 个结果。",
        "显示候选阈值为 0.35，机械动作阈值为 0.45，并要求同类连续 3 个处理周期。",
        "界面和机械控制只采用最终得分最高的一个目标，避免多目标同时出现时路由不确定。",
    ])
    add_callout(doc, "准确率原则", "识别不准时先补真实部署视角的数据、核对标注和训练/验证划分，再看混淆矩阵。不要把阈值一降到底，否则低分候选会直接增加机械误动作。", ORANGE)

    add_heading(doc, "12. SDRAM 与内存占用", 2)
    add_image(doc, GEN / "memory_usage.png", 17.0, "图5 当前双核内存分区和构建占用")
    add_p(doc, "CPU0 构建：text 434,490 B、data 108 B、bss 995,264 B；CPU1 构建：text 9,470 B、data 8 B、bss 1,408 B。按当前 Solution 分区计算，两个核心均未溢出。外部 SDRAM 已完成初始化和读写验证，但目前不能仅凭地址可见性宣称 NPU 工作区已经直接放入 SDRAM。")

    add_heading(doc, "13. 拍照写卡与现场数据闭环", 2)
    add_numbers(doc, [
        "进入 DATASET 模式后选择 HARM、KITCHEN、RECOVER 或 OTHER。",
        "程序从当前完整 CEU 帧按模型同一路径生成 192×192 图像。",
        "图像编码为 24-bit BGR BMP，通过 FatFs 写入对应类别目录。",
        "文件名使用递增序号，启动时扫描已有文件，避免覆盖。",
        "单拍用于控制样本质量；自动采集用于快速积累姿态、背景和光照变化。",
    ])
    add_callout(doc, "为什么存储仍在 M85", "写卡直接使用刚完成采集的帧。若迁到 M33，需要共享 SDRAM、缓存一致性、跨核所有权和 SDHI DMA 协调，复制成本和风险高于当前收益。因此先迁移触摸、称重和健康监测，存储保留在 M85。", BLUE)

    chapter(doc, "第四篇 程序与机械闭环")
    add_heading(doc, "14. 关键代码所有权", 2)
    add_table(doc, ["文件/模块", "核心", "职责"], [
        ["src/hal_entry.c", "M85", "初始化、主循环、预览、推理、状态与心跳"],
        ["src/ov5640.c", "M85", "SCCB、复位、寄存器和图像格式"],
        ["src/yolo_detector.c", "M85 + U55", "预处理、模型调用、双尺度解码、NMS、平滑"],
        ["src/machine_control.c", "M85", "步进、舵机、视觉对齐与分拣状态机"],
        ["src/dataset_storage.c", "M85", "FatFs 挂载、目录、BMP 编码和写卡"],
        ["src/platform_services.c", "M85", "IPC 快照、掉线降级、心跳发送"],
        ["multicore/m33/m33_services.c", "M33", "触摸、称重、健康包与心跳监测"],
        ["src/multicore_protocol.h", "两核共享", "32 位 IPC 消息格式"],
    ], widths=[5.2, 2.8, 9.0])

    add_heading(doc, "15. 上电与主循环", 2)
    add_numbers(doc, [
        "初始化显示，使系统尽早具备可见诊断能力。",
        "初始化平台服务和 IPC，M85 启动 M33；M33 初始化触摸与称重。",
        "复位并配置 OV5640，读取 ID，打开 CEU。",
        "打开 Ethos-U55 驱动，初始化模型、机械控制和 SD 卡服务。",
        "启动第一帧采集，循环处理完整帧、执行 NPU、更新最高目标、推进状态机和刷新显示。",
        "M33 独立轮询触摸与称重；任何服务核故障不应永久阻塞视觉主链路。",
    ])
    if (ASSETS / "image4.png").exists():
        add_image(doc, ASSETS / "image4.png", 16.2, "图6 系统软件总体流程与异常恢复")

    add_heading(doc, "16. 分拣状态机", 2)
    add_table(doc, ["状态", "动作", "退出条件"], [
        ["WAITING", "正转送料，等待最高目标稳定", "置信度 ≥0.45 且同类连续 3 周期"],
        ["ALIGNING", "根据框中心与判定线偏差低速正/反转", "中心落入 160±12 px；丢框 6 周期或超时也继续"],
        ["POSITIONING", "按类别向前或向后输送", "前 10,000 ms；后 4,000 ms，需继续实机标定"],
        ["EJECTING", "连续舵机向目标方向推料", "1,800 ms"],
        ["CLOSING", "舵机反向旋转回位", "1,800 ms 后输出停转脉宽"],
        ["RETURNING", "等待机构稳定", "250 ms"],
        ["CLEARING", "传送带恢复正转", "画面无目标 4 周期后重新识别"],
    ], widths=[3.0, 8.0, 6.0])
    if (ASSETS / "image11.png").exists():
        add_image(doc, ASSETS / "image11.png", 16.5, "图7 识别结果到机械动作的非阻塞状态机")
    add_callout(doc, "防重复机制", "类别一旦稳定识别就锁定。目标在返回判定线时即使暂时丢失，ALIGN 阶段连续丢框 6 次后也从当前位置继续，不退回 WAITING 反复识别同一电池。", BLUE)

    add_heading(doc, "17. 触摸界面与模式", 2)
    add_table(doc, ["按钮/模式", "功能"], [
        ["START / RUN", "启用视觉推理和自动分拣"],
        ["PAUSE", "保留预览，暂停机构动作，便于放置样品"],
        ["STOP / END", "停止自动流程并让执行器进入安全状态"],
        ["DATA", "进入数据集采集界面"],
        ["H/K/R/O", "选择有害、厨余、可回收、其他类别"],
        ["蓝色拍照", "保存当前 192×192 BMP；自动模式按设定间隔连续保存"],
    ], widths=[4.5, 12.5])

    chapter(doc, "第五篇 FSP、构建、烧录与验收")
    add_heading(doc, "18. 双核 Solution 当前配置", 2)
    add_table(doc, ["项目", "关键外设", "内存"], [
        ["rainy_dual_CPU0", "CEU、Ethos-U、SPI1、SDHI0、DMAC、GPT、IPC0", "Flash 0xE0000；RAM 0x180000；SDRAM 128 MB"],
        ["rainy_dual_CPU1", "IIC1、SCI4、IPC0、GPIO", "Flash 0x20000；RAM 0x54000"],
    ], widths=[4.5, 7.0, 5.5])
    add_bullets(doc, [
        "CPU0 的 IPC 实例名为 g_ipc_m85，CPU1 为 g_ipc_m33；两边 Channel 都是 0。",
        "CPU1 IIC1：P511=SDA、P512=SCL、P105=触摸复位。",
        "CPU1 SCI4：P715=RX、P714=TX，9600 8N1，SCISPICLK/SCICLK 必须启用。",
        "CPU0 不再占用 IIC1 和 SCI4；CPU1 不配置摄像头、显示、SDHI、NPU 和电机。",
        "CPU0 宏 PLATFORM_SERVICES_MULTICORE=1；单核回退时可改为 0。",
    ])

    add_heading(doc, "19. e² studio 编译与烧录", 2)
    add_numbers(doc, [
        "在 Solution Configuration 中保存 Memories、Clocks 和两核外设归属。",
        "分别打开 CPU0、CPU1 的 configuration.xml，点击 Generate Project Content。",
        "先 Clean，再 Build 两个工程；确认 CPU0/CPU1 均为 0 errors。",
        "选择 rainy_dual_CPU1 Debug_Multicore Launch Group，一次下载两个 ELF。",
        "若弹出 active configuration 有错误，不要直接忽略；先查看 Problems。若只是已知配置提示且两个工程已正常 Build，才继续。",
        "调试器目标频率警告可在 Debug Tool Settings 中设置实际 Operating Frequency；它不等同于下载失败。",
        "下载后复位，观察状态块、摄像头预览、触摸、称重和 IPC 健康。",
    ])
    add_callout(doc, "下载成功的证据", "CPU1 日志中 Target connection status - OK 只说明已连接目标；还要看到 Finished download，并在板上确认 M33 健康包到达。", ORANGE)

    add_heading(doc, "20. 建议验收矩阵", 2)
    add_table(doc, ["测试", "方法", "通过标准"], [
        ["冷启动", "连续断电上电 5 次", "均进入预览，无随机红屏/花屏"],
        ["双核 IPC", "运行 10 分钟并观察状态", "M33 健康包持续；固定时间心跳改完后不再黄绿交替"],
        ["触摸", "RUN/PAUSE/STOP/DATA 各 20 次", "无漏按、误按，不影响摄像头画面"],
        ["称重", "零点和标准砝码", "数据刷新，倍率与零点可解释"],
        ["SD 写卡", "单拍 100 张 + 自动采集 10 分钟", "文件可在电脑逐张读回，无覆盖和卡死"],
        ["四类分拣", "每类至少 10 次", "方向、到位、推出、回位、清场正确"],
        ["掉线降级", "暂停 CPU1 超过 3 秒", "M33 标红；视觉、显示与安全状态不被锁死"],
    ], widths=[3.1, 6.2, 7.7])

    add_heading(doc, "21. 常见故障排查", 2)
    add_table(doc, ["现象", "优先检查"], [
        ["彩条正常、实景花", "像素格式、字节交换、PCLK/D0～D7 信号完整性"],
        ["一帧正常一帧花", "双缓冲交换、CEU 未结束时写屏/写卡、缓存一致性"],
        ["插触摸后画面花", "P511/P512/P105 线序、共地回流、I2C 轮询时机和线束干扰"],
        ["左半红", "CPU1 是否下载和启动、IPC Channel/IRQ/callback"],
        ["左半黄绿交替", "M85 心跳按 32 帧发送；改为固定 500 ms 时间心跳"],
        ["右半红", "FT6336 地址 0x38、RST、SDA/SCL 电平和上拉"],
        ["SD CARD ERR", "卡格式 FAT32、挂载、SDHI/DMAC、写入阶段错误码"],
        ["识别后机构重复循环", "类别锁定、ALIGN 丢框计数、CLEAR 无目标计数"],
    ], widths=[5.0, 12.0])

    chapter(doc, "第六篇 三分钟讲稿、问答与交接")
    add_heading(doc, "22. 三分钟完整讲稿", 2)
    script = [
        ("0:00–0:25", "各位老师好，我们的作品是基于瑞萨 RA8P1 的轻量化 AI 视觉感知与智能分拣平台。垃圾分类只是一个可观察载体，我们真正探索的是一颗 MCU 如何同时完成视觉、AI、交互、存储和实时控制。"),
        ("0:25–0:55", "系统由 OV5640、ILI9488 触摸屏、称重、TF 卡和步进舵机构成。CEU 获取 320×240 图像，Cortex-M85 将图像处理为 192×192，Ethos-U55 完成 int8 离线推理。"),
        ("0:55–1:25", "当前版本进一步使用 Cortex-M33 协同。M33 独立负责触摸、称重和健康监测，M85 专注视觉推理和整机状态机，两核通过 32 位 IPC 传输小数据，不跨核搬运整帧。"),
        ("1:25–1:55", "控制端只采用最高置信度目标，并通过显示阈值、执行阈值、连续帧确认和类别锁定抑制误动作。目标先在判定线附近闭环对齐，再输送到对应舵机完成推出、反向回位和清场。"),
        ("1:55–2:25", "外部 SDRAM 已完成初始化和读写验证；TF 卡拍照写入也已跑通，设备可以保存与模型输入一致的 192×192 BMP，形成现场难例采集和再训练的数据闭环。"),
        ("2:25–3:00", "因此，作品的价值不只在四类垃圾识别，而在于验证 RA8P1 的 M85、M33、Ethos-U55、CEU、SDHI、SDRAM、SPI、I2C、UART 和 GPT 如何协同构成完整边缘智能系统。更换模型和机械末端后，同一平台可以迁移到零件、农产品和缺陷检测。谢谢各位老师。"),
    ]
    add_table(doc, ["时间", "讲稿"], script, widths=[2.7, 14.3], font_size=9.5)

    add_heading(doc, "23. 高频答辩问答", 2)
    qa = [
        ("为什么要用 M33？", "不是因为 M85 算力不够，而是把 I2C、UART 和健康监测隔离成服务核，让视觉主循环更确定，也真实使用 RA8P1 双核资源。"),
        ("双核之间传图像吗？", "当前不传。IPC 只传 32 位触摸、重量、健康和心跳消息；大图像仍由 M85 侧帧缓冲处理。"),
        ("M33 掉线会怎样？", "M85 超时后把服务核标红，触摸和称重降级，但视觉、显示和执行安全状态不应被 IPC 永久阻塞。"),
        ("为什么 SD 卡没有放 M33？", "拍照写卡直接依赖完整帧。迁移会引入共享 SDRAM、缓存一致性和 SDHI DMA 所有权，当前收益不够。"),
        ("NPU 已经使用 SDRAM 吗？", "SDRAM 已可用，但目前没有完成 NPU 工作区迁移验证，所以只能说平台具备条件，不能说已经直接使用。"),
        ("为什么输入 192×192？", "它在计算量、小目标信息和内存之间折中，适合轻量 YOLO 和当前闭环帧率。"),
        ("为什么相机仍是 320×240？", "QVGA 为预览、框定位和判定线保留空间，模型前再缩放到 192×192。"),
        ("为什么只处理一个目标？", "机构一次只分拣一个物体，最高置信度目标让控制对象唯一，再用连续帧确认降低误触发。"),
        ("显示阈值和执行阈值为什么不同？", "显示 0.35 方便观察候选，执行 0.45 且连续 3 周期，防止低分候选直接驱动机构。"),
        ("连续舵机为什么要反向回位？", "连续舵机 PWM 表示速度和方向，不表示绝对角度；停止脉宽只能停住，所以必须反转相同时长回位。"),
        ("为什么不用 delay 控制动作？", "长 delay 会冻结相机、触摸、显示和故障恢复。状态机每次循环只推进一步。"),
        ("是否把 RA8P1 所有功能都用上了？", "没有，也没有必要。我们重点使用了视觉、AI、双核、存储、交互和控制相关资源，并明确保留 MIPI、GLCDC、OSPI、安全和网络扩展方向。"),
    ]
    for q, a in qa:
        add_p(doc, "Q  " + q, bold=True, color=BLUE, size=10.5, after=1, keep=True)
        add_p(doc, "A  " + a, size=10.2, after=5)

    add_heading(doc, "24. 队友一天接手路线", 2)
    add_numbers(doc, [
        "先背熟项目定位、30 秒介绍和三分钟讲稿，不要先陷入寄存器细节。",
        "看图4理解 M85/M33/U55 分工，再看图1理解从像素到动作的数据流。",
        "在代码中依次阅读 hal_entry、platform_services、m33_services、yolo_detector、machine_control。",
        "对着附录引脚表在实物上指出摄像头、屏幕、触摸、称重、舵机和步进接口。",
        "完整演示 RUN、PAUSE、STOP、DATA、四类识别和一次 SD 拍照。",
        "练习回答第23章问题，并能区分“已验证”“已集成待长测”“后续方向”。",
    ])

    chapter(doc, "附录A 当前参数与文件速查")
    add_heading(doc, "A.1 关键参数", 2)
    add_table(doc, ["参数", "当前值", "用途"], [
        ["相机帧", "320×240 RGB565", "预览、检测和数据采集来源"],
        ["模型输入", "192×192 int8", "Ethos-U55 推理"],
        ["显示/执行阈值", "0.35 / 0.45", "观察候选 / 触发机构"],
        ["稳定确认", "3 周期", "同类连续出现"],
        ["判定线", "Y=160，容差 ±12 px", "视觉对齐"],
        ["前/后输送", "10,000 / 4,000 ms", "到前/后舵机位置"],
        ["推料/回位", "1,800 / 1,800 ms", "连续舵机动作"],
        ["M33 健康/超时", "1,000 / 3,000 ms", "服务核在线判断"],
    ], widths=[4.0, 5.0, 8.0])
    add_heading(doc, "A.2 工程证据", 2)
    add_bullets(doc, [
        "D:\\RA\\second\\rainy\\src\\platform_services.c：M85 IPC 服务、超时和心跳。",
        "D:\\RA\\second\\rainy\\multicore\\m33\\m33_services.c：M33 触摸、称重和健康循环。",
        "D:\\RA\\second\\rainy\\src\\multicore_protocol.h：32 位消息协议。",
        "D:\\RA\\second\\rainy\\src\\yolo_detector.c：192×192、NMS、平滑和候选阈值。",
        "D:\\RA\\second\\rainy\\src\\machine_control.c：分拣路线、计时和防重复。",
        "D:\\RA\\second\\rainy_dual_CPU0\\Debug\\rainy_dual_CPU0.elf：M85 镜像。",
        "D:\\RA\\second\\rainy_dual_CPU1\\Debug\\rainy_dual_CPU1.elf：M33 镜像。",
    ])
    add_heading(doc, "A.3 比赛前最终清单", 2)
    add_bullets(doc, [
        "比赛前确认：两核 Clean/Build 为 0 errors，复合下载后 M33 健康包可见。",
        "比赛前待完成：将 M85 心跳改为固定 500 ms 时间驱动，并验证左半状态不再因低帧率黄绿交替。",
        "比赛前待验证：冷启动 5 次、触摸 80 次、单拍 100 张、自动写卡 10 分钟，并记录结果。",
        "比赛前待验证：四类各至少 10 次，记录成功率、误检、漏检和机械失败原因。",
        "准备 SD 卡目录、BMP 样本、双核状态视频、SDRAM 测试和 Build 资源截图。",
        "三位队员都能说明已实现边界，不把待验证功能当成成绩。",
    ])

    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / "RA8P1异构边缘智能与智能分拣平台_答辩全景手册_当前双核版.docx"
    doc.save(path)
    return path


def build_comparison():
    doc = Document()
    setup_document(doc, "RA8P1 当前实现与原文档差异对照｜2026-08-20")
    cover(doc, "RA8P1 当前实现与原文档差异对照", "文档修订、证据与剩余验证项", "用于报告修订、答辩口径统一与版本审查")
    add_heading(doc, "1. 修订结论", 1)
    add_callout(doc, "当前版本已经发生的根本变化", "项目已从 M85 单核集中式闭环，推进到 M85 视觉主核 + M33 外设服务核的双核架构；SDRAM 与 SD 拍照写入也已从“规划/待联调”变为已实现基础能力。", BLUE)
    add_p(doc, "本对照不用于三分钟 PPT 讲旧方案，而是帮助队员修改报告、回答版本问题，并确保所有陈述都有工程或实机证据。")

    add_heading(doc, "2. 功能逐项对照", 1)
    rows = [
        ["处理器架构", "M85 单核，M33 作为后续方向", "M85/M33 双核 Solution 已建立并编译", "已实现；继续长时间实机复测"],
        ["触摸", "M85 直接轮询 FT6336", "M33 IIC1 轮询，IPC 发送触摸", "已迁移并编译"],
        ["称重", "M85 SCI4 接收", "M33 SCI4 接收并通过 IPC 上报", "已迁移并编译；复测实机刷新"],
        ["健康监测", "局部状态灯", "M33 健康包、M85/M33 心跳、3秒超时降级", "已实现；心跳需改固定时间"],
        ["TF/SD", "待联调或部分完成", "192×192 BMP 拍照写卡已能工作", "已实现；补 100 张和 10 分钟稳定性"],
        ["SDRAM", "未使用或未验证", "0x68000000、128 MB，初始化和读写可用", "已验证基础读写"],
        ["NPU 与 SDRAM", "可能混写为 NPU 已使用", "尚未完成 NPU arena 迁移验证", "不得过度宣称"],
        ["构建资源", "单核 1 MiB Flash 口径", "CPU0 896 KB + CPU1 128 KB 分区", "两核均未溢出"],
        ["烧录", "单 ELF 下载", "Multicore Launch Group 同时下载两个 ELF", "配置已建立"],
    ]
    add_table(doc, ["项目", "旧文档口径", "当前口径", "状态"], rows, widths=[2.8, 4.5, 6.0, 3.7], font_size=8.8)

    add_heading(doc, "3. 必须替换的旧表述", 1)
    add_table(doc, ["不要再写", "改为"], [
        ["“当前没有使用 Cortex-M33。”", "“M33 已承担触摸、称重和健康监测，M85 专注视觉、NPU、存储和执行控制。”"],
        ["“SDRAM 尚未使用/未验证。”", "“SDRAM 已完成初始化和读写验证，当前 128 MB 归 CPU0；NPU 工作区迁移仍待验证。”"],
        ["“SD 卡拍照待联调。”", "“数据集模式可保存 192×192 BMP；赛前继续做批量读回和长时间写入测试。”"],
        ["“双核使用 IPC0 与 IPC1 两个通道。”", "“两核各一个 IPC 实例，名称不同，但 Channel 均为 0。”"],
        ["“先烧 M33 再烧 M85 即可。”", "“优先使用 rainy_dual_CPU1 Debug_Multicore Launch Group 同时下载两个 ELF。”"],
        ["“左上角黄色一定是故障。”", "“黄色表示 M33 在线但尚未确认 M85 心跳；当前按帧心跳可能导致黄绿交替。”"],
    ], widths=[8.0, 9.0])

    add_heading(doc, "4. 当前证据矩阵", 1)
    add_table(doc, ["结论", "工程证据", "现场证据"], [
        ["CPU1 可编译", "rainy_dual_CPU1.elf；text 9470、data 8、bss 1408", "Build Finished, 0 errors, 0 warnings"],
        ["CPU0 可编译", "rainy_dual_CPU0.elf；text 434490、data 108、bss 995264", "Build Finished, 0 errors"],
        ["双核协议已实现", "multicore_protocol.h、platform_services.c、m33_services.c", "状态块能收到 M33 健康包"],
        ["SDRAM 可用", "Solution SDRAM 0x68000000 / 0x08000000", "读写测试已通过"],
        ["SD 写卡可用", "dataset_storage + SDHI/DMAC/FatFs", "电脑可看到 192×192 BMP"],
        ["NPU 视觉可用", "yolo_detector + 转换生成模型", "实时画面有框、类别和置信度"],
        ["机械闭环可用", "machine_control 状态机", "推出、反转回位和清场流程"],
    ], widths=[3.5, 7.5, 6.0])

    add_heading(doc, "5. 当前资源口径", 1)
    add_image(doc, GEN / "memory_usage.png", 17.0, "当前双核构建与内存分区")
    add_bullets(doc, [
        "CPU0 Flash 使用约 47.37%，内部 RAM 静态使用约 63.29%。",
        "CPU1 Flash 使用约 7.23%，内部 RAM 静态使用约 0.41%。",
        "外部 SDRAM 128 MB 当前全部归 CPU0，避免第一阶段跨核缓存一致性风险。",
        "ELF 文件大小不等于 Flash 实际占用，答辩应引用 llvm-size 和分区容量。",
    ])

    add_heading(doc, "6. 仍然不能当作完成项的内容", 1)
    add_table(doc, ["事项", "为什么不能过度宣称", "下一步证据"], [
        ["NPU 直接访问 SDRAM", "尚未验证模型 arena/命令流的实际放置和缓存一致性", "map、section、地址与推理稳定性"],
        ["双核长期稳定", "当前已构建和启动，但心跳仍受帧率影响", "改为 500 ms 心跳，连续运行 30 分钟"],
        ["M33 存储服务", "SDHI 当前由 M85 持有，且直接依赖帧缓冲", "共享缓冲协议和缓存维护实验"],
        ["最终识别准确率", "尚无统一测试集和混淆矩阵", "冻结测试集、每类指标和现场成功率"],
        ["机械成功率", "行程和舵机时间仍需按实机标定", "每类至少 10～30 次记录"],
    ], widths=[4.0, 7.5, 5.5])

    add_heading(doc, "7. 材料更新规则", 1)
    add_numbers(doc, [
        "三分钟 PPT 只讲当前平台，不讲旧硬件方案，也不放版本比较。",
        "全景手册以当前双核版为唯一主文档；原补充说明只保留归档价值。",
        "本对照作为报告修订附录，所有“已完成”都必须能对应代码、Build 或实机证据。",
        "赛前每次功能变化都同步更新：状态表、构建资源、演示步骤和高频问答。",
        "无法稳定复现的功能写成“已集成待复测”，不要为了答辩把边界说满。",
    ])

    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / "RA8P1当前实现与原文档差异对照_2026-08-20.docx"
    doc.save(path)
    return path


def main():
    make_diagrams()
    main_path = build_main()
    compare_path = build_comparison()
    print(main_path)
    print(compare_path)


if __name__ == "__main__":
    main()
